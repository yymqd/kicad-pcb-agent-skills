#!/usr/bin/env python3
"""Generate clean customer-facing design report (.docx).
Call after Phase 9 design review. No FMEA, no dev history."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, shutil

DOCX = "设计报告.docx"
IMG = "exports"

# ══ REPORT CONTENT ══

def build():
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('设计报告')
    r.font.size = Pt(22); r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('图号：_____    版本：_____    日期：_____').font.size = Pt(10)
    doc.add_page_break()

    # ══ 1. 电路原理 ══
    doc.add_heading('一、电路原理', level=1)
    doc.add_heading('1.1 拓扑结构', level=2)
    doc.add_paragraph('描述电路拓扑、架构框图、工作原理。说明输入→Boost→恒流源→负载的功率流向。')
    doc.add_heading('1.2 功能模块', level=2)
    # Fill module rows from caller
    t = doc.add_table(rows=4, cols=3); t.style = 'Light Shading Accent 1'
    for i, d in enumerate([('模块','核心器件','功能'),('Boost','U1 + L1 + D1','升压至...'),('恒流源','U2 + Q1 + RS','I=Vref/Rs'),('控制/保护','...','...')]):
        for j, txt in enumerate(d): t.rows[i].cells[j].text = txt

    # ══ 2. 设计参数 ══
    doc.add_heading('二、设计参数', level=1)
    t = doc.add_table(rows=7, cols=3); t.style = 'Light Shading Accent 1'
    for i, d in enumerate([('参数','值','说明'),('输入电压','',''),('Boost输出','',''),
                           ('Vref','',''),('恒流输出','',''),('调制','',''),('PCB','','')]):
        for j, txt in enumerate(d): t.rows[i].cells[j].text = txt

    # ══ 3. PCB 布局 ══
    doc.add_heading('三、PCB 布局', level=1)
    doc.add_paragraph('尺寸、层数、布局分区说明。')
    for fname in ['pcb_view.png', 'pcb_3d.png']:
        fp = os.path.join(IMG, fname)
        if os.path.exists(fp):
            doc.add_picture(fp, width=Inches(4.5))
            doc.paragraphs[-1].alignment = 1

    doc.add_heading('布局分区', level=2)
    t = doc.add_table(rows=5, cols=3); t.style = 'Light Shading Accent 1'
    for i, d in enumerate([('区域','位置','元件'),('电源','',''),('控制','',''),('输出','',''),('保护','','')]):
        for j, txt in enumerate(d): t.rows[i].cells[j].text = txt

    # ══ 4. BOM ══
    doc.add_heading('四、物料清单（BOM）', level=1)
    bom = [('Ref','型号/值','封装','数量')]
    # Add rows from caller
    t = doc.add_table(rows=len(bom), cols=4); t.style = 'Light Shading Accent 1'
    for i, row in enumerate(bom):
        for j, txt in enumerate(row): t.rows[i].cells[j].text = txt

    # ══ 5. 送厂打样 ══
    doc.add_heading('五、送厂打样', level=1)
    t = doc.add_table(rows=7, cols=2); t.style = 'Light Shading Accent 1'
    for i, (k, v) in enumerate([('参数','值'),('层数','2'),('板厚','1.6mm'),('铜厚','1 oz'),
                                 ('表面处理','HASL'),('阻焊颜色','绿色'),('上传文件','Gerbers.zip')]):
        t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v
    doc.add_paragraph('上传 Gerber ZIP 到 JLCPCB 下单。')

    # ══ 6. 文件清单 ══
    doc.add_heading('六、文件清单', level=1)
    t = doc.add_table(rows=7, cols=3); t.style = 'Light Shading Accent 1'
    for i, d in enumerate([('文件','大小','说明'),('board.kicad_pcb','','KiCad源文件'),
                           ('Gerbers.zip','','打样文件'),('design_report.docx','','本报告'),
                           ('pcb_view.png','','布局预览'),('pcb_3d.png','','3D渲染图'),
                           ('positions.csv','','贴片坐标')]):
        for j, txt in enumerate(d): t.rows[i].cells[j].text = txt

    # ══ 7. 测试验证 ══
    doc.add_heading('七、测试验证', level=1)
    t = doc.add_table(rows=4, cols=3); t.style = 'Light Shading Accent 1'
    for i, d in enumerate([('测试项','方法','预期'),('测试1','',''),('测试2','',''),('测试3','','')]):
        for j, txt in enumerate(d): t.rows[i].cells[j].text = txt

    # ══ 8. 仿真验证 ══
    doc.add_heading('八、仿真验证', level=1)
    doc.add_paragraph('仿真方法和关键结果。')
    t = doc.add_table(rows=5, cols=3); t.style = 'Light Shading Accent 1'
    for i, d in enumerate([('参数','仿真值','期望值'),('','',''),('','',''),
                           ('','',''),('','','')]):
        for j, txt in enumerate(d): t.rows[i].cells[j].text = txt

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('— 设计自动生成 by pcb-design skill —').font.size = Pt(9)

    doc.save(DOCX)
    print(f"✅ {DOCX} ({os.path.getsize(DOCX)//1024} KB)")

if __name__ == '__main__':
    build()
